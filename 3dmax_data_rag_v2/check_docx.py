import docx


def check_document(file_path):
    """Simple function to check if we can open and read the document."""
    try:
        # Try to open the document
        doc = docx.Document(file_path)

        # Print information about the document
        print(f"Successfully opened document: {file_path}")
        print(f"Number of paragraphs: {len(doc.paragraphs)}")

        # Print the first few paragraphs to see what's in the document
        print("\nFirst 5 paragraphs content:")
        for i, para in enumerate(doc.paragraphs[:5]):
            print(f"Paragraph {i + 1}: {para.text[:100]}...")

            # Print run information for debugging
            print(f"  Paragraph {i + 1} has {len(para.runs)} runs")
            for j, run in enumerate(para.runs[:3]):  # First 3 runs only
                print(f"    Run {j + 1}: '{run.text[:20]}...' (Bold: {run.bold}, Italic: {run.italic})")

    except Exception as e:
        print(f"Error opening document: {e}")


if __name__ == "__main__":
    file_path = "data/test_doc.docx"
    check_document(file_path)