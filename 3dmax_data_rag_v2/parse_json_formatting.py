import re
import json
from docx import Document
import os
from docx.text.paragraph import Paragraph


def extract_media_references(text):
    """Extract media references (image files) from text"""
    # Pattern matches references like image1.gif, image15.png, etc.
    media_refs = []
    for match in re.finditer(r'image\d+\.(gif|png|jpg|jpeg)', text):
        media_refs.append(match.group(0))

    return list(set(media_refs))  # Remove duplicates


def convert_formatting_to_html(paragraph):
    """Convert DOCX paragraph formatting to HTML"""
    html_parts = []

    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Escape HTML special characters
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # Apply formatting
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"

        html_parts.append(text)

    return "".join(html_parts)


def process_document(document_path):
    # Load the document
    document = Document(document_path)
    print(f"Document loaded with {len(document.paragraphs)} paragraphs")

    # First, convert each paragraph to HTML with formatting
    formatted_paragraphs = []
    for para in document.paragraphs:
        if para.text.strip():
            formatted_text = convert_formatting_to_html(para)
            formatted_paragraphs.append(formatted_text)

    # Join all paragraphs with newlines to create a single formatted text
    full_formatted_text = "\n".join(formatted_paragraphs)

    # Now we'll extract questions and answers from the formatted text
    # Initialize our data structure
    result = {
        "sections": [
            {
                "id": 1,
                "title": "Основное руководство (единый алгоритм решения большинства задач)",
                "content": []
            },
            {
                "id": 2,
                "title": "Вспомогательные решения",
                "content": []
            },
            {
                "id": 3,
                "title": "База знаний: 230 вопросов и ответов",
                "subcategories": []
            }
        ]
    }

    # Define subcategories for section 3
    subcategories = [
        {"letter": "A", "title": "Общие вопросы по интерфейсу", "range": (1, 10)},
        {"letter": "B", "title": "Моделирование", "range": (11, 40)},
        {"letter": "C", "title": "Анимация", "range": (41, 60)},
        {"letter": "D", "title": "Материалы и текстуры", "range": (61, 80)},
        {"letter": "E", "title": "Освещение", "range": (81, 100)},
        {"letter": "F", "title": "Рендеринг", "range": (101, 120)},
        {"letter": "G", "title": "Скриптинг и автоматизация", "range": (121, 140)},
        {"letter": "H", "title": "Динамика и симуляция", "range": (141, 160)},
        {"letter": "I", "title": "Импорт/Экспорт", "range": (161, 170)},
        {"letter": "J", "title": "Плагины и интеграция", "range": (171, 180)},
        {"letter": "K", "title": "Оптимизация и устранение ошибок", "range": (181, 190)},
        {"letter": "L", "title": "Дополнительная оптимизация", "range": (191, 200)},
        {"letter": "M", "title": "Продвинутые техники UV-развертки", "range": (201, 210)},
        {"letter": "N", "title": "Продвинутые материалы", "range": (211, 220)},
        {"letter": "O", "title": "Продвинутая анимация и работа с камерой", "range": (221, 230)},
        {"letter": "P", "title": "Баги и ошибки", "range": (231, 250)}
    ]

    # Create subcategories in the result
    for subcat in subcategories:
        letter = subcat["letter"]
        title = subcat["title"]
        start_q, end_q = subcat["range"]

        result["sections"][2]["subcategories"].append({
            "id": f"3.{letter}",
            "title": f"{letter}. {title} (Вопросы {start_q}--{end_q})",
            "content": []
        })

    # Extract all questions and their positions in the text
    questions = []
    # Pattern to find questions with HTML formatting
    pattern = r'<strong>Вопрос (\d+)</strong>(?:<strong>:</strong>|:)\s*(.*?)(?=<strong>Вопрос \d+</strong>|<strong>Ответ:</strong>|$)'

    # Simplify processing by using a plain text version alongside formatted text
    # This helps with position matching
    plain_text = full_formatted_text
    for tag in ['<strong>', '</strong>', '<em>', '</em>']:
        plain_text = plain_text.replace(tag, '')

    # Extract questions directly from the document text
    for para in document.paragraphs:
        text = para.text.strip()
        # Look for question pattern
        match = re.search(r'Вопрос (\d+):', text)
        if match:
            q_num = int(match.group(1))
            # Extract question text (everything after "Вопрос X:")
            q_text = text[text.find(':') + 1:].strip()
            questions.append((q_num, q_text))
            print(f"Found Question {q_num}: {q_text[:50]}...")

    # For each question, find its answer
    for q_num, q_text in questions:
        # Find the corresponding answer in the document
        answer_text = ""
        answer_found = False

        # Get the answer by looking through paragraphs
        for i in range(len(document.paragraphs)):
            para = document.paragraphs[i].text.strip()
            if not para:
                continue

            # Check if this paragraph contains the question
            if f"Вопрос {q_num}:" in para:
                # Look for the answer in subsequent paragraphs
                j = i + 1
                while j < len(document.paragraphs):
                    next_para = document.paragraphs[j].text.strip()
                    j += 1

                    # Skip empty paragraphs
                    if not next_para:
                        continue

                    # Stop if we hit the next question
                    if re.search(r'Вопрос \d+:', next_para):
                        break

                    # If this paragraph contains the answer marker, we found it
                    if "Ответ:" in next_para:
                        answer_found = True
                        # Remove the "Ответ:" prefix
                        answer_text = next_para.replace("Ответ:", "").strip()

                        # Continue collecting answer paragraphs
                        k = j
                        while k < len(document.paragraphs):
                            next_answer_para = document.paragraphs[k].text.strip()
                            k += 1

                            # Stop if we hit the next question
                            if re.search(r'Вопрос \d+:', next_answer_para):
                                break

                            # Add this paragraph to the answer
                            if next_answer_para:
                                answer_text += "\n" + next_answer_para

                        break

                # If we found the answer, no need to check more paragraphs
                if answer_found:
                    break

        # Apply formatting to question and answer text
        formatted_question = q_text
        formatted_answer = answer_text

        # Convert explicit formatting markers to HTML tags
        # This is a simplified approach - a more robust HTML conversion would be better
        patterns = [
            (r'\[\*\*([^*]+)\*\*\]', r'<strong>\1</strong>'),  # [**text**]
            (r'\*\*([^*]+)\*\*', r'<strong>\1</strong>'),  # **text**
            (r'\[([^\]]+)\]', r'<strong>\1</strong>'),  # [text]
            (r'\{\.mark\}', '')  # {.mark}
        ]

        # Apply patterns to question and answer
        for pattern, replacement in patterns:
            formatted_question = re.sub(pattern, replacement, formatted_question)
            formatted_answer = re.sub(pattern, replacement, formatted_answer)

        # Extract media references from the answer
        media_refs = extract_media_references(formatted_answer)

        # Find the subcategory for this question
        for i, subcat in enumerate(subcategories):
            start_q, end_q = subcat["range"]
            if start_q <= q_num <= end_q:
                letter = subcat["letter"]

                # Create the QA item
                qa_item = {
                    "id": f"3.{letter}.{q_num}",
                    "question": formatted_question,
                    "answer": formatted_answer,
                    "media": media_refs
                }

                # Add to the appropriate subcategory
                result["sections"][2]["subcategories"][i]["content"].append(qa_item)
                print(f"Added Q{q_num} to subcategory {letter}")
                break

    # Sort questions within each subcategory
    for subcat in result["sections"][2]["subcategories"]:
        subcat["content"].sort(key=lambda x: int(x["id"].split(".")[-1]))

    return result


def main():
    docx_path = "data_short.docx"  # Path to your DOCX file
    output_path = "3ds_max_qa_formatting.json"  # Output JSON file

    # Verify the document exists
    if not os.path.exists(docx_path):
        print(f"Error: Document not found at {docx_path}")
        return

    # Process the document
    print(f"Processing document: {docx_path}")
    result = process_document(docx_path)

    # Check content
    section3 = result["sections"][2]
    subcategories = section3.get("subcategories", [])
    print(f"Found {len(subcategories)} subcategories in section 3")

    total_qa_pairs = 0
    for subcat in subcategories:
        qa_count = len(subcat.get("content", []))
        total_qa_pairs += qa_count
        print(f"Subcategory {subcat['id']} has {qa_count} Q&A pairs")

    print(f"Total Q&A pairs extracted: {total_qa_pairs}")

    # Write the result to JSON
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Conversion complete! Output saved to {output_path}")


if __name__ == "__main__":
    main()