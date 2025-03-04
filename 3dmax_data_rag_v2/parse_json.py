import re
import json
from docx import Document
import os


def extract_media_references(text):
    """Extract media references (image files) from text"""
    # Pattern matches references like image1.gif, image15.png, etc.
    media_refs = []
    for match in re.finditer(r'image\d+\.(gif|png|jpg|jpeg)', text):
        media_refs.append(match.group(0))

    return list(set(media_refs))  # Remove duplicates


def clean_text(text):
    """Clean text while preserving formatting as HTML"""
    # Convert strong formatting
    text = re.sub(r'\[\*\*([^*]+)\*\*\]', r'<strong>\1</strong>', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\[([^\]]+)\]', r'<strong>\1</strong>', text)

    # Remove {.mark} tags
    text = re.sub(r'\{\.mark\}', '', text)

    # Remove multiple spaces and normalize line breaks
    text = re.sub(r'\s+', ' ', text).strip()

    return text


def extract_qa_pairs(document):
    """Extract question-answer pairs from the document"""
    # Initialize result structure
    result = {
        "sections": [
            {
                "id": 1,
                "title": "Основное руководство (единый алгоритм решения большинства задач)",
                "content": []
            },
            {
                "id": 2,
                "title": "Вспомогательные решения",
                "content": []
            },
            {
                "id": 3,
                "title": "База знаний: 230 вопросов и ответов",
                "subcategories": []
            }
        ]
    }

    # Define subcategories and their question ranges manually
    subcategories_info = [
        {"letter": "A", "title": "Общие вопросы по интерфейсу", "range": (1, 10)},
        {"letter": "B", "title": "Моделирование", "range": (11, 40)},
        {"letter": "C", "title": "Анимация", "range": (41, 60)},
        {"letter": "D", "title": "Материалы и текстуры", "range": (61, 80)},
        {"letter": "E", "title": "Освещение", "range": (81, 100)},
        {"letter": "F", "title": "Рендеринг", "range": (101, 120)},
        {"letter": "G", "title": "Скриптинг и автоматизация", "range": (121, 140)},
        {"letter": "H", "title": "Динамика и симуляция", "range": (141, 160)},
        {"letter": "I", "title": "Импорт/Экспорт", "range": (161, 170)},
        {"letter": "J", "title": "Плагины и интеграция", "range": (171, 180)},
        {"letter": "K", "title": "Оптимизация и устранение ошибок", "range": (181, 190)},
        {"letter": "L", "title": "Дополнительная оптимизация", "range": (191, 200)},
        {"letter": "M", "title": "Продвинутые техники UV-развертки", "range": (201, 210)},
        {"letter": "N", "title": "Продвинутые материалы", "range": (211, 220)},
        {"letter": "O", "title": "Продвинутая анимация и работа с камерой", "range": (221, 230)},
        {"letter": "P", "title": "Баги и ошибки", "range": (231, 250)}
    ]

    # Create subcategories in the result
    for subcat_info in subcategories_info:
        letter = subcat_info["letter"]
        title = subcat_info["title"]
        start_q, end_q = subcat_info["range"]

        subcategory = {
            "id": f"3.{letter}",
            "title": f"{letter}. {title} (Вопросы {start_q}--{end_q})",
            "content": []
        }

        result["sections"][2]["subcategories"].append(subcategory)
        print(f"Created subcategory {letter}: {title} (Questions {start_q}-{end_q})")

    # Now let's extract all paragraphs and work with them directly
    all_text = ""
    for para in document.paragraphs:
        if para.text.strip():
            all_text += para.text.strip() + "\n"

    # Split the text into more manageable chunks
    chunks = all_text.split("\n")

    # Find all questions and their positions
    questions = []
    for i, chunk in enumerate(chunks):
        # Search for "Вопрос X:" pattern in various formats
        matches = re.finditer(r"(?:\*\*)?(?:Вопрос|Bomрос)\s+(\d+)(?::|\.|\))(?:\*\*)?\s*([^\n]*)", chunk)
        for match in matches:
            q_num = int(match.group(1))
            q_text = match.group(2).strip()
            questions.append((q_num, q_text, i))
            print(f"Found Question {q_num} at position {i}: {q_text[:50]}...")

    # Sort questions by position in the document
    questions.sort(key=lambda x: x[2])

    # Process each question and find its answer
    qa_pairs = []
    for i, (q_num, q_text, q_pos) in enumerate(questions):
        # The answer should be between this question and the next
        answer_text = ""

        # Find where the answer starts
        answer_start = q_pos
        for j in range(q_pos, len(chunks)):
            if "Ответ:" in chunks[j]:
                answer_start = j
                break

        # Determine where the answer ends (at the next question or the end of the document)
        answer_end = len(chunks)
        if i < len(questions) - 1:
            answer_end = questions[i + 1][2]

        # Extract the answer text
        answer_chunks = chunks[answer_start:answer_end]
        if answer_chunks:
            # Remove the "Ответ:" prefix from the first chunk
            if "Ответ:" in answer_chunks[0]:
                answer_chunks[0] = answer_chunks[0].split("Ответ:", 1)[1].strip()

            answer_text = " ".join(answer_chunks).strip()

        # Extract media references
        media_refs = extract_media_references(answer_text)

        # Store the Q&A pair
        qa_pairs.append({
            "number": q_num,
            "question": clean_text(q_text),
            "answer": clean_text(answer_text),
            "media": media_refs
        })

        print(f"Processed Q{q_num}: Answer has {len(media_refs)} media references")

    # Now distribute the Q&A pairs to the appropriate subcategories
    for qa_pair in qa_pairs:
        q_num = qa_pair["number"]

        # Find the right subcategory
        for i, subcat_info in enumerate(subcategories_info):
            start_q, end_q = subcat_info["range"]
            if start_q <= q_num <= end_q:
                letter = subcat_info["letter"]

                # Create the final Q&A item
                qa_item = {
                    "id": f"3.{letter}.{q_num}",
                    "question": qa_pair["question"],
                    "answer": qa_pair["answer"],
                    "media": qa_pair["media"]
                }

                # Add to the correct subcategory
                result["sections"][2]["subcategories"][i]["content"].append(qa_item)
                print(f"Added Q{q_num} to subcategory {letter}")
                break

    # Sort questions within each subcategory by number
    for subcat in result["sections"][2]["subcategories"]:
        subcat["content"].sort(key=lambda x: int(x["id"].split(".")[-1]))

    return result


def convert_docx_to_json(docx_path):
    """Convert the DOCX file to JSON structure"""
    try:
        # Load the document
        document = Document(docx_path)
        print(f"Document loaded with {len(document.paragraphs)} paragraphs")

        # Extract question-answer pairs
        result = extract_qa_pairs(document)

        # Check content
        section3 = result["sections"][2]
        subcategories = section3.get("subcategories", [])
        print(f"Found {len(subcategories)} subcategories in section 3")

        total_qa_pairs = 0
        for subcat in subcategories:
            qa_count = len(subcat.get("content", []))
            total_qa_pairs += qa_count
            print(f"Subcategory {subcat['id']} has {qa_count} Q&A pairs")

        print(f"Total Q&A pairs extracted: {total_qa_pairs}")

        return result

    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def main():
    docx_path = "data_complete.docx"  # Path to your DOCX file
    output_path = "3ds_max_qa.json"  # Output JSON file

    # Verify the document exists
    if not os.path.exists(docx_path):
        print(f"Error: Document not found at {docx_path}")
        return

    print(f"Reading document: {docx_path}")

    # Convert DOCX to JSON
    result = convert_docx_to_json(docx_path)

    if result:
        # Write the JSON output
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"Conversion complete! Output saved to {output_path}")
    else:
        print("Conversion failed!")


if __name__ == "__main__":
    main()