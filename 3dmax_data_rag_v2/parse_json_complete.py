import re
import json
from docx import Document
import os


def extract_media_references(text):
    """Extract media references (image files) from text"""
    # Pattern matches references like image1.gif, image15.png, etc.
    media_refs = []
    for match in re.finditer(r'image\d+\.(gif|png|jpg|jpeg)', text):
        media_refs.append(match.group(0))

    return list(set(media_refs))  # Remove duplicates


def clean_text(text):
    """Clean text while preserving formatting as HTML"""
    # Convert strong formatting
    text = re.sub(r'\[\*\*([^*]+)\*\*\]', r'<strong>\1</strong>', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\[([^\]]+)\]', r'<strong>\1</strong>', text)

    # Remove {.mark} tags
    text = re.sub(r'\{\.mark\}', '', text)

    # Remove multiple spaces and normalize line breaks
    text = re.sub(r'\s+', ' ', text).strip()

    return text


def extract_all_numbered_items(document):
    """Extract all numbered items from the document"""
    # Initialize result structure
    result = {
        "sections": [
            {
                "id": 1,
                "title": "Основное руководство (единый алгоритм решения большинства задач)",
                "content": []
            }
        ]
    }

    # Create a combined text string of all paragraphs
    # But keep track of paragraph numbers for debugging
    paragraphs = []
    for i, para in enumerate(document.paragraphs):
        if para.text.strip():
            paragraphs.append((i, para.text.strip()))

    # Print the first few paragraphs for debugging
    print("First 5 paragraphs:")
    for i, (para_num, text) in enumerate(paragraphs[:5]):
        print(f"Paragraph {para_num}: {text[:100]}...")

    # Now extract all items that appear to be numbered with content
    current_item = None
    current_title = ""
    current_content = []

    for para_num, text in paragraphs:
        # Try different patterns for numbered items
        # Pattern 1: Digit followed by a period and a title in bold
        match1 = re.match(r'(\d+)\.\s+\*\*([^*]+)\*\*', text)
        # Pattern 2: Digit followed by period, then title with formatting
        match2 = re.match(r'(\d+)\.\s+([A-Za-zА-Яа-я].+?):', text)

        if match1 or match2:
            # If we were collecting content for a previous item, save it
            if current_item is not None:
                joined_content = " ".join(current_content)
                media_refs = extract_media_references(joined_content)

                # Only add items with ID 1-5 (assuming these are Section 1 items)
                if 1 <= current_item <= 5:
                    result["sections"][0]["content"].append({
                        "id": f"1.{current_item}",
                        "title": current_title,
                        "content": clean_text(joined_content),
                        "media": media_refs
                    })
                    print(f"Added item {current_item}: '{current_title}' with {len(media_refs)} media references")

            # Start a new item
            match = match1 or match2
            current_item = int(match.group(1))
            current_title = match.group(2).strip().replace('**', '')

            # If the item title is at the beginning of the paragraph,
            # the rest of the paragraph is the start of the content
            rest_of_para = text
            if match1:
                title_pos = text.find(f"**{current_title}**")
                if title_pos != -1:
                    rest_of_para = text[title_pos + len(f"**{current_title}**"):].strip()
                    if rest_of_para.startswith(':'):
                        rest_of_para = rest_of_para[1:].strip()
            elif match2:
                title_pos = text.find(f"{current_title}:")
                if title_pos != -1:
                    rest_of_para = text[title_pos + len(f"{current_title}:"):].strip()

            current_content = [rest_of_para] if rest_of_para else []

            print(f"Found item {current_item} at paragraph {para_num}: '{current_title}'")
        elif current_item is not None:
            # Continue collecting content for current item
            current_content.append(text)

    # Don't forget to save the last item
    if current_item is not None:
        joined_content = " ".join(current_content)
        media_refs = extract_media_references(joined_content)

        # Only add items with ID 1-5 (assuming these are Section 1 items)
        if 1 <= current_item <= 5:
            result["sections"][0]["content"].append({
                "id": f"1.{current_item}",
                "title": current_title,
                "content": clean_text(joined_content),
                "media": media_refs
            })
            print(f"Added item {current_item}: '{current_title}' with {len(media_refs)} media references")

    # Sort items by ID
    result["sections"][0]["content"].sort(key=lambda x: int(x["id"].split(".")[1]))

    return result


def main():
    docx_path = "data_complete.docx"  # Path to your DOCX file
    output_path = "section1_content.json"  # Output JSON file

    # Verify the document exists
    if not os.path.exists(docx_path):
        print(f"Error: Document not found at {docx_path}")
        return

    print(f"Reading document: {docx_path}")

    try:
        # Load the document
        document = Document(docx_path)
        print(f"Document loaded with {len(document.paragraphs)} paragraphs")

        # Extract all numbered items
        result = extract_all_numbered_items(document)

        # Count extracted items
        section1_items = len(result["sections"][0]["content"])
        print(f"\nExtraction summary:")
        print(f"Section 1: {section1_items} items extracted")

        # Write the JSON output
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"Conversion complete! Output saved to {output_path}")

    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()